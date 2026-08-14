"""Optional Microsoft Word DOCX renderer.

The renderer always builds a new package.  It does not import relationships,
macros, OLE objects, metadata, or externally referenced images from a source
document.
"""

from __future__ import annotations

import math
import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from ..errors import RenderError
from ..model import (
    Annotation,
    CharacterStyle,
    Footer,
    Footnote,
    Frame,
    Header,
    Image,
    PageBreak,
    Paragraph,
    SdwDrawing,
    StyleDefinition,
    Table,
    TableCell,
    TableRow,
    TextRun,
    UnsupportedObject,
    WmfGraphic,
    _paragraph_text,
    _TextOutputBudget,
)
from ..model import Document as AmiProDocument
from ..sdw import SdwDecodeError, sdw_display_size, sdw_png, sdw_preview_caption
from ..wmf import WmfDecodeError, wmf_display_size, wmf_png

__all__ = ["render"]


_HEX_COLOR = re.compile(r"#?([0-9a-fA-F]{6})\Z")
_RSID_ATTRIBUTE = re.compile(rb'\s+w:rsid[A-Za-z0-9]*="[^"]*"')
_RSID_SETTINGS = re.compile(rb"<w:rsids(?:\s[^>]*)?>.*?</w:rsids\s*>", re.DOTALL)
_DOC_ID = re.compile(rb"<w14:docId(?:\s[^>]*)?/>\s*")
_SAVE_PREVIEW = re.compile(rb"<w:savePreviewPicture(?:\s[^>]*)?/>\s*")
_ZIP_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
_MAX_TABLE_COLUMNS = 256
_MAX_TABLE_ROWS = 390
_COVERED = object()
_MAX_PAGE_TWIPS = 31_680
_MIN_PAGE_TWIPS = 1_440
_MIN_BODY_TWIPS = 720
_MIN_FURNITURE_MARGIN_TWIPS = 720
_MAX_BLOCKS_PER_LIST = 100_000
_MAX_BLOCK_DEPTH = 64
_INVALID_BLOCK_CONTAINER = object()
_BLOCK_LIMIT_OMISSION = object()


@dataclass(frozen=True, slots=True)
class _PageGeometry:
    width_twips: int = 12_240
    height_twips: int = 15_840
    margin_left_twips: int = 1_440
    margin_right_twips: int = 1_440
    margin_top_twips: int = 1_440
    margin_bottom_twips: int = 1_440
    layout_index: int | None = None
    non_alternating: bool = False

    @property
    def body_width_twips(self) -> int:
        return self.width_twips - self.margin_left_twips - self.margin_right_twips

    @property
    def body_height_twips(self) -> int:
        return self.height_twips - self.margin_top_twips - self.margin_bottom_twips


@dataclass(frozen=True, slots=True)
class _NativePageContent:
    header_odd: Header | None = None
    header_even: Header | None = None
    footer_odd: Footer | None = None
    footer_even: Footer | None = None
    ids: frozenset[int] = frozenset()
    alternating: bool = True

    @property
    def uses_odd_even(self) -> bool:
        return self.alternating and any(
            item is not None
            for item in (
                self.header_odd,
                self.header_even,
                self.footer_odd,
                self.footer_even,
            )
        )


def _page_geometry(document: object) -> _PageGeometry:
    layouts = getattr(document, "page_layouts", None)
    if not isinstance(layouts, (list, tuple)):
        return _PageGeometry()
    for layout in layouts:
        if getattr(layout, "valid", None) is not True:
            continue
        for variant_name in ("odd", "even"):
            variant = getattr(layout, variant_name, None)
            if getattr(variant, "valid", None) is not True:
                continue
            values = tuple(
                getattr(variant, name, None)
                for name in (
                    "width_twips",
                    "height_twips",
                    "margin_left_twips",
                    "margin_right_twips",
                    "margin_top_twips",
                    "margin_bottom_twips",
                )
            )
            if not all(type(value) is int for value in values):
                continue
            width, height, left, right, top, bottom = values
            if not (
                _MIN_PAGE_TWIPS <= width <= _MAX_PAGE_TWIPS
                and _MIN_PAGE_TWIPS <= height <= _MAX_PAGE_TWIPS
            ):
                continue
            if min(left, right, top, bottom) < 0:
                continue
            if left + right >= width or top + bottom >= height:
                continue
            if width - left - right < _MIN_BODY_TWIPS:
                continue
            if height - top - bottom < _MIN_BODY_TWIPS:
                continue
            layout_index = getattr(layout, "index", None)
            if type(layout_index) is not int:
                layout_index = None
            return _PageGeometry(
                width_twips=width,
                height_twips=height,
                margin_left_twips=left,
                margin_right_twips=right,
                margin_top_twips=top,
                margin_bottom_twips=bottom,
                layout_index=layout_index,
                non_alternating=getattr(layout, "non_alternating", None) is True,
            )
    return _PageGeometry()


def _native_page_content(
    document: object,
    geometry: _PageGeometry,
) -> _NativePageContent:
    if geometry.layout_index is None:
        return _NativePageContent()
    candidates: dict[str, list[Header | Footer]] = {
        "header_odd": [],
        "header_even": [],
        "footer_odd": [],
        "footer_even": [],
    }
    for block in _safe_blocks(getattr(document, "blocks", None)):
        if not isinstance(block, Header | Footer):
            continue
        placement = getattr(block, "placement", None)
        if not isinstance(placement, str) or placement not in {"odd", "even"}:
            continue
        if getattr(block, "origin", None) != "layout":
            continue
        if getattr(block, "layout_index", None) != geometry.layout_index:
            continue
        key = f"{'header' if isinstance(block, Header) else 'footer'}_{placement}"
        candidates[key].append(block)
    selected = {
        key: (
            values[0]
            if len(values) == 1 and _native_furniture_fits(values[0], geometry)
            else None
        )
        for key, values in candidates.items()
    }
    if geometry.non_alternating:
        selected["header_even"] = None
        selected["footer_even"] = None
    ids = frozenset(id(value) for value in selected.values() if value is not None)
    return _NativePageContent(
        ids=ids,
        alternating=not geometry.non_alternating,
        **selected,
    )


def _native_furniture_fits(
    block: Header | Footer,
    geometry: _PageGeometry,
) -> bool:
    if getattr(block, "terminated", None) is not True:
        return False
    if type(getattr(block, "unknown_flag_bits", None)) is not int:
        return False
    if block.unknown_flag_bits != 0:
        return False
    frame = getattr(block, "frame", None)
    if frame is not None and (
        not isinstance(frame, Frame)
        or type(getattr(frame, "unknown_flag_bits", None)) is not int
        or frame.unknown_flag_bits != 0
    ):
        return False
    child_blocks = getattr(block, "blocks", None)
    if not isinstance(child_blocks, (list, tuple)) or not 1 <= len(child_blocks) <= 4:
        return False
    if not all(_native_paragraph_is_safe(item) for item in child_blocks):
        return False
    characters_per_line = max(20, geometry.body_width_twips // 144)
    line_count = 0
    total_characters = 0
    for paragraph in child_blocks:
        text = _bounded_native_paragraph_text(paragraph)
        if text is None:
            return False
        total_characters += len(text)
        if total_characters > 8_192:
            return False
        lines = text.splitlines() or [""]
        line_count += sum(
            max(1, math.ceil(len(line) / characters_per_line)) for line in lines
        )
    margin = (
        geometry.margin_top_twips
        if isinstance(block, Header)
        else geometry.margin_bottom_twips
    )
    required = max(_MIN_FURNITURE_MARGIN_TWIPS, 360 + line_count * 360)
    return margin >= required


def _bounded_native_paragraph_text(paragraph: Paragraph) -> str | None:
    runs = getattr(paragraph, "runs", None)
    if not isinstance(runs, (list, tuple)) or len(runs) > 1_024:
        return None
    parts: list[str] = []
    length = 0
    for run in runs:
        if not isinstance(run, TextRun) or not isinstance(run.style, CharacterStyle):
            return None
        value = run.text
        if not isinstance(value, str):
            return None
        length += len(value)
        if length > 4_096:
            return None
        parts.append(value)
    return "".join(parts)


def _native_paragraph_is_safe(value: object) -> bool:
    page_break = getattr(value, "page_break_before", None)
    list_kind = getattr(value, "list_kind", None)
    return (
        isinstance(value, Paragraph)
        and (page_break is False or page_break is None)
        and (list_kind is None or (isinstance(list_kind, str) and list_kind == ""))
        and isinstance(getattr(value, "runs", None), (list, tuple))
    )


def _safe_blocks(value: object) -> list[object]:
    if not isinstance(value, (list, tuple)):
        return [_INVALID_BLOCK_CONTAINER]
    result: list[object] = list(value[:_MAX_BLOCKS_PER_LIST])
    if len(value) > _MAX_BLOCKS_PER_LIST:
        result.append(_BLOCK_LIMIT_OMISSION)
    return result


def _normalized_blocks(blocks: list[object]) -> list[object]:
    """Discard edge breaks while preserving every explicit interior break."""

    result: list[object] = []
    index = 0
    while index < len(blocks):
        if not isinstance(blocks[index], PageBreak):
            result.append(blocks[index])
            index += 1
            continue
        end = index
        while end < len(blocks) and isinstance(blocks[end], PageBreak):
            end += 1
        if result and end < len(blocks):
            result.extend(blocks[index:end])
        index = end
    return result


def _frame_label(frame: Frame) -> str:
    content_kind = getattr(frame, "content_kind", None)
    if not isinstance(content_kind, str) or content_kind not in {
        "text", "table", "image", "drawing", "unknown"
    }:
        content_kind = "unknown"
    placement = getattr(frame, "placement", None)
    if not isinstance(placement, str) or placement not in {
        "anchored", "fixed-page", "repeating", "unknown"
    }:
        placement = "unknown"
    return (
        f"[Frame: {content_kind}; {placement} placement reflowed in source order; "
        "original coordinates not reproduced]"
    )


def render(document: AmiProDocument, **_options: object) -> bytes:
    """Return *document* as DOCX bytes.

    ``python-docx`` is an optional dependency so importing this module remains
    cheap.  Selecting DOCX without the extra installed produces a direct,
    actionable error rather than an import traceback.
    """

    try:
        from docx import Document as WordDocument
        from docx.enum.section import WD_ORIENT
        from docx.shared import Pt, Twips
    except ImportError as exc:
        raise RenderError(
            "DOCX output requires the optional 'python-docx' dependency. "
            "Install it with `pip install 'amipro-sam-toolkit[docx]'` "
            "(or `pip install python-docx`)."
        ) from exc

    try:
        geometry = _page_geometry(document)
        native = _native_page_content(document, geometry)
        output_document = WordDocument()
        section = output_document.sections[0]
        section.orientation = (
            WD_ORIENT.LANDSCAPE
            if geometry.width_twips > geometry.height_twips
            else WD_ORIENT.PORTRAIT
        )
        section.page_width = Twips(geometry.width_twips)
        section.page_height = Twips(geometry.height_twips)
        section.top_margin = Twips(geometry.margin_top_twips)
        section.right_margin = Twips(geometry.margin_right_twips)
        section.bottom_margin = Twips(geometry.margin_bottom_twips)
        section.left_margin = Twips(geometry.margin_left_twips)
        section.header_distance = Twips(min(708, geometry.margin_top_twips // 2))
        section.footer_distance = Twips(min(708, geometry.margin_bottom_twips // 2))

        normal = output_document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1
        _scrub_core_properties(output_document.core_properties)

        text_budget = _TextOutputBudget()
        _add_native_page_content(output_document, section, native, text_budget)
        _add_blocks(
            output_document,
            document,
            document.blocks,
            geometry,
            native.ids,
            active_container_ids=set(),
            seen_block_ids=set(),
            text_budget=text_budget,
        )

        buffer = BytesIO()
        output_document.save(buffer)
        return _sanitize_package(buffer.getvalue())
    except RenderError:
        raise
    except Exception as exc:
        raise RenderError(f"Could not render DOCX safely: {exc}") from exc


def _add_blocks(
    target: Any,
    document: AmiProDocument,
    blocks: object,
    geometry: _PageGeometry,
    native_ids: frozenset[int],
    *,
    depth: int = 0,
    active_container_ids: set[int],
    seen_block_ids: set[int],
    text_budget: _TextOutputBudget,
) -> None:
    if depth > _MAX_BLOCK_DEPTH:
        _add_placeholder(target, "[Nested content omitted: safe depth limit reached]")
        return
    if isinstance(blocks, (list, tuple)):
        container_id = id(blocks)
        if container_id in active_container_ids:
            _add_placeholder(
                target,
                "[Nested content omitted: repeated or cyclic block reference]",
            )
            return
        active_container_ids.add(container_id)
    else:
        container_id = None
    visible = [
        block for block in _safe_blocks(blocks) if id(block) not in native_ids
    ]
    try:
        for block in _normalized_blocks(visible):
            block_identity = id(block)
            if block_identity in seen_block_ids:
                _add_placeholder(
                    target,
                    "[Nested content omitted: repeated or cyclic block reference]",
                )
                continue
            seen_block_ids.add(block_identity)
            if isinstance(block, Paragraph):
                paragraph = target.add_paragraph()
                _populate_paragraph(paragraph, block, document, text_budget)
            elif isinstance(block, PageBreak):
                target.add_page_break()
            elif isinstance(block, Table):
                _add_table(
                    target,
                    block,
                    document,
                    geometry.body_width_twips,
                    text_budget,
                )
            elif isinstance(block, Image):
                _add_placeholder(target, _image_placeholder(block))
            elif isinstance(block, WmfGraphic):
                _add_wmf(target, block, geometry)
            elif isinstance(block, SdwDrawing):
                _add_sdw(target, block, geometry)
            elif isinstance(block, Frame):
                _add_placeholder(target, _frame_label(block))
                _add_blocks(
                    target,
                    document,
                    getattr(block, "blocks", None),
                    geometry,
                    native_ids,
                    depth=depth + 1,
                    active_container_ids=active_container_ids,
                    seen_block_ids=seen_block_ids,
                    text_budget=text_budget,
                )
            elif isinstance(block, UnsupportedObject):
                kind = _safe_label_field(
                    block.kind, "unknown object kind", maximum=128
                )
                description = _safe_label_field(
                    block.description, "description unavailable", maximum=256
                )
                _add_placeholder(target, f"[Unsupported {kind}: {description}]")
            elif isinstance(block, Annotation | Footnote | Header | Footer):
                _add_placeholder(target, _container_label(block))
                _add_blocks(
                    target,
                    document,
                    getattr(block, "blocks", None),
                    geometry,
                    native_ids,
                    depth=depth + 1,
                    active_container_ids=active_container_ids,
                    seen_block_ids=seen_block_ids,
                    text_budget=text_budget,
                )
            elif block is _INVALID_BLOCK_CONTAINER:
                _add_placeholder(target, "[Invalid block container omitted]")
            elif block is _BLOCK_LIMIT_OMISSION:
                _add_placeholder(
                    target, "[Block content omitted at safe rendering limit]"
                )
            else:
                _add_placeholder(target, "[Unrecognized block object omitted]")
    finally:
        # Keep visited containers for the full render to bound shared DAGs.
        pass


def _add_native_page_content(
    output_document: Any,
    section: Any,
    native: _NativePageContent,
    text_budget: _TextOutputBudget,
) -> None:
    if native.uses_odd_even:
        output_document.settings.odd_and_even_pages_header_footer = True
    for container, source in (
        (section.header if native.header_odd is not None else None, native.header_odd),
        (
            section.even_page_header if native.header_even is not None else None,
            native.header_even,
        ),
        (section.footer if native.footer_odd is not None else None, native.footer_odd),
        (
            section.even_page_footer if native.footer_even is not None else None,
            native.footer_even,
        ),
    ):
        if container is None or source is None:
            continue
        container.is_linked_to_previous = False
        _populate_native_container(
            container,
            getattr(source, "blocks", None),
            text_budget,
        )


def _populate_native_container(
    container: Any,
    blocks: object,
    text_budget: _TextOutputBudget,
) -> None:
    paragraphs = _safe_blocks(blocks)
    if not paragraphs:
        return
    existing = container.paragraphs
    for index, source in enumerate(paragraphs):
        target = existing[0] if index == 0 and existing else container.add_paragraph()
        _clear_word_paragraph(target)
        if isinstance(source, Paragraph):
            target.add_run(
                _clean_xml_text(
                    _paragraph_text(source, text_budget, expansion_factor=5)
                )
            )


def _clear_word_paragraph(paragraph: Any) -> None:
    properties = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not properties:
            paragraph._p.remove(child)


def _container_label(block: Annotation | Footnote | Header | Footer) -> str:
    if isinstance(block, Annotation):
        return "[Annotation]"
    if isinstance(block, Footnote):
        return (
            "[Footnote "
            + _safe_label_field(
                block.number, "number unavailable", maximum=64
            )
            + "]"
            if block.number is not None
            else "[Footnote]"
        )
    kind = "Header" if isinstance(block, Header) else "Footer"
    source_placement = getattr(block, "placement", None)
    if not isinstance(source_placement, str):
        source_placement = "unknown"
    placement = {
        "all": "all pages",
        "odd": "odd/right pages",
        "even": "even/left pages",
        "odd-even": "odd and even variants",
        "unknown": "placement unknown",
    }.get(source_placement, "placement unknown")
    return f"[{kind}: {placement}]"


def _add_wmf(target: Any, graphic: WmfGraphic, geometry: _PageGeometry) -> None:
    from docx.shared import Inches

    try:
        payload = wmf_png(graphic)
        width, height = wmf_display_size(
            graphic,
            max_width_in=min(6.25, geometry.body_width_twips / 1440),
            max_height_in=min(7.5, geometry.body_height_twips / 1440),
        )
    except WmfDecodeError:
        _add_placeholder(target, "[Invalid WMF preview]")
        return
    paragraph = target.add_paragraph()
    paragraph.add_run(
        _clean_xml_text(
            _safe_label_field(
                graphic.alt_text, "Embedded WMF preview", maximum=256
            )
        )
    )
    paragraph.add_run().add_picture(
        BytesIO(payload), width=Inches(width), height=Inches(height)
    )


def _add_sdw(target: Any, drawing: SdwDrawing, geometry: _PageGeometry) -> None:
    from docx.shared import Inches

    try:
        payload = sdw_png(drawing)
        width, height = sdw_display_size(
            drawing,
            max_width_in=min(6.25, geometry.body_width_twips / 1440),
            max_height_in=min(7.5, geometry.body_height_twips / 1440),
        )
    except SdwDecodeError:
        _add_placeholder(target, _sdw_placeholder(drawing))
        return
    if (
        not isinstance(payload, bytes)
        or not payload.startswith(b"\x89PNG\r\n\x1a\n")
        or not _valid_sdw_display_size(width, height)
    ):
        _add_placeholder(target, _sdw_placeholder(drawing))
        return
    label = target.add_paragraph()
    label.add_run(sdw_preview_caption(drawing))
    paragraph = target.add_paragraph()
    inline_shape = paragraph.add_run().add_picture(
        BytesIO(payload), width=Inches(width), height=Inches(height)
    )
    alt = _clean_xml_text(
        _safe_label_field(drawing.alt_text, "Ami Draw object", maximum=256)
    )
    inline_shape._inline.docPr.set("descr", alt)
    inline_shape._inline.docPr.set("title", alt)


def _populate_paragraph(
    target: Any,
    source: Paragraph,
    document: AmiProDocument,
    text_budget: _TextOutputBudget,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    style_definition = _resolved_style(document, source.style_name)
    if source.list_kind:
        level = _safe_level(source.list_level)
        style_suffix = "" if level == 0 else f" {min(level + 1, 3)}"
        style_name = f"List {'Number' if source.list_kind == 'number' else 'Bullet'}{style_suffix}"
        try:
            target.style = style_name
        except KeyError:
            target.style = f"List {'Number' if source.list_kind == 'number' else 'Bullet'}"

    alignment = source.alignment or (style_definition.alignment if style_definition else None)
    target.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(alignment)
    formatting = target.paragraph_format
    left = _first_not_none(
        source.left_indent_in,
        style_definition.left_indent_in if style_definition else None,
    )
    if left is None and source.list_kind:
        left = 0.5 + 0.25 * _safe_level(source.list_level)
    right = _first_not_none(
        source.right_indent_in,
        style_definition.right_indent_in if style_definition else None,
    )
    first = _first_not_none(
        source.first_line_indent_in,
        style_definition.first_line_indent_in if style_definition else None,
    )
    before = _first_not_none(
        source.space_before_pt,
        style_definition.space_before_pt if style_definition else None,
    )
    after = _first_not_none(
        source.space_after_pt,
        style_definition.space_after_pt if style_definition else None,
    )
    spacing = _first_not_none(
        source.line_spacing,
        style_definition.line_spacing if style_definition else None,
    )
    if left is not None:
        formatting.left_indent = Inches(_number(left, 0.0, -20.0, 20.0))
    if right is not None:
        formatting.right_indent = Inches(_number(right, 0.0, -20.0, 20.0))
    if first is not None:
        formatting.first_line_indent = Inches(_number(first, 0.0, -20.0, 20.0))
    if before is not None:
        formatting.space_before = Pt(_number(before, 0.0, 0.0, 720.0))
    if after is not None:
        formatting.space_after = Pt(_number(after, 0.0, 0.0, 720.0))
    if spacing is not None:
        formatting.line_spacing = _number(spacing, 1.2, 0.5, 10.0)
    formatting.page_break_before = bool(source.page_break_before)
    formatting.keep_with_next = bool(source.keep_with_next)

    base = style_definition.character if style_definition else CharacterStyle()
    source_runs = source.runs
    if not isinstance(source_runs, list | tuple):
        target.add_run("[Invalid paragraph runs omitted]")
        return
    seen_runs: set[int] = set()
    omitted = len(source_runs) > 4_096
    total_characters = 0
    for source_run in source_runs[:4_096]:
        if not isinstance(source_run, TextRun) or not isinstance(
            source_run.style, CharacterStyle
        ):
            target.add_run("[Invalid text run omitted]")
            continue
        if id(source_run) in seen_runs:
            omitted = True
            continue
        seen_runs.add(id(source_run))
        value = source_run.text
        if isinstance(value, bytes):
            value = value.decode("utf-8", errors="replace")
        if not isinstance(value, str):
            target.add_run("[Invalid text run omitted]")
            continue
        paragraph_remaining = max(0, 1_000_000 - total_characters)
        prepared = text_budget.prepare(
            value,
            unit_limit=paragraph_remaining,
            expansion_factor=5,
        )
        value = prepared.visible
        total_characters += len(prepared.text)
        if prepared.encoding in {"bounded-text", "text-budget-limit"}:
            omitted = True
        run = target.add_run(_clean_xml_text(value))
        _format_run(run, _merge_character_style(base, source_run.style))
        if total_characters >= 1_000_000:
            omitted = True
            break
    if omitted:
        target.add_run("[Paragraph content omitted at safe rendering limit]")


def _format_run(run: Any, style: CharacterStyle) -> None:
    from docx.enum.text import WD_UNDERLINE
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.bold = bool(style.bold)
    run.italic = bool(style.italic)
    run.underline = WD_UNDERLINE.SINGLE if style.underline else False
    run.font.strike = bool(style.strike)
    run.font.superscript = bool(style.superscript)
    run.font.subscript = bool(style.subscript and not style.superscript)
    if isinstance(style.font_family, str) and style.font_family:
        family = _clean_xml_text(style.font_family)[:128]
        run.font.name = family
        run_properties = run._element.get_or_add_rPr()
        fonts = run_properties.get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "cs", "eastAsia"):
            fonts.set(qn(f"w:{attribute}"), family)
    if style.font_size_pt is not None:
        run.font.size = Pt(_number(style.font_size_pt, 11.0, 1.0, 200.0))
    color = _color(style.color)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_placeholder(document: Any, text: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(_clean_xml_text(text))
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F4F4")
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "B8B8B8")
    borders.append(bottom)
    properties.append(borders)


def _add_table(
    document: Any,
    source: Table,
    ir_document: AmiProDocument,
    body_width_twips: int,
    text_budget: _TextOutputBudget,
) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

    rows = _safe_table_rows(source)
    if not rows:
        _add_placeholder(
            document,
            "[Invalid table rows omitted]"
            if not isinstance(source.rows, list | tuple)
            else "[Empty table]",
        )
        return
    if any(
        isinstance(row.cells, list | tuple)
        and len(row.cells) > _MAX_TABLE_COLUMNS
        for row in rows
    ):
        _add_placeholder(
            document,
            "[Table cells omitted at safe 256-column limit]",
        )
        return
    try:
        grid, anchors = _layout_table(source)
    except RenderError:
        _add_placeholder(document, "[Table grid omitted at safe 256-column limit]")
        return
    if not grid or not grid[0]:
        _add_placeholder(document, "[Empty table]")
        return
    row_count = len(grid)
    column_count = len(grid[0])
    table = document.add_table(rows=row_count, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = _configure_table_geometry(table, column_count, body_width_twips)

    merged_cells: dict[tuple[int, int], Any] = {}
    for row_index, column_index, _cell, column_span, row_span in anchors:
        target = table.cell(row_index, column_index)
        if column_span > 1 or row_span > 1:
            target = target.merge(
                table.cell(row_index + row_span - 1, column_index + column_span - 1)
            )
            _set_cell_width(target, sum(widths[column_index : column_index + column_span]))
        merged_cells[(row_index, column_index)] = target

    seen_cells: set[int] = set()
    for row_index, column_index, source_cell, _column_span, _row_span in anchors:
        target = merged_cells[(row_index, column_index)]
        target.text = ""
        target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        repeated_cell = id(source_cell) in seen_cells
        seen_cells.add(id(source_cell))
        paragraphs = (
            source_cell.blocks[:4_096]
            if isinstance(source_cell.blocks, list | tuple)
            else []
        )
        valid_paragraphs: list[Paragraph] = []
        seen_paragraphs: set[int] = set()
        for item in paragraphs:
            if not isinstance(item, Paragraph) or id(item) in seen_paragraphs:
                continue
            seen_paragraphs.add(id(item))
            valid_paragraphs.append(item)
        invalid_content = repeated_cell or (
            not isinstance(source_cell.blocks, list | tuple)
            or len(valid_paragraphs) != len(paragraphs)
            or (
                isinstance(source_cell.blocks, list | tuple)
                and len(source_cell.blocks) > 4_096
            )
        )
        if invalid_content:
            target.paragraphs[0].add_run(
                "[Repeated table cell omitted]"
                if repeated_cell
                else "[Invalid or repeated table cell content omitted]"
            )
        for block_index, source_paragraph in enumerate(
            valid_paragraphs if not repeated_cell else []
        ):
            paragraph_index = block_index + int(invalid_content)
            paragraph = (
                target.paragraphs[0]
                if paragraph_index == 0
                else target.add_paragraph()
            )
            _populate_paragraph(
                paragraph,
                source_paragraph,
                ir_document,
                text_budget,
            )
        if rows[row_index].is_header:
            _shade_cell(target, "F2F4F7")
            for paragraph in target.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    for row_index, row in enumerate(rows):
        if row.is_header:
            _shade_header_row(table.rows[row_index])
        if row_index < 8 and row.is_header and all(
            previous.is_header for previous in rows[: row_index + 1]
        ):
            _set_repeat_table_header(table.rows[row_index])


def _configure_table_geometry(
    table: Any,
    column_count: int,
    body_width_twips: int,
) -> list[int]:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Twips

    total_width = max(_MIN_BODY_TWIPS, body_width_twips)
    base, remainder = divmod(total_width, column_count)
    widths = [base + (1 if index < remainder else 0) for index in range(column_count)]
    table_properties = table._tbl.tblPr
    _set_or_add_measure(table_properties, "w:tblW", total_width, "dxa")
    _set_or_add_measure(table_properties, "w:tblInd", 0, "dxa")
    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    cell_margins = table_properties.find(qn("w:tblCellMar"))
    if cell_margins is None:
        cell_margins = OxmlElement("w:tblCellMar")
        table_properties.append(cell_margins)
    for side, width in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        _set_or_add_measure(cell_margins, f"w:{side}", width, "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            _set_cell_width(cell, widths[index])
    return widths


def _set_cell_width(cell: Any, width: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    _set_or_add_measure(properties, "w:tcW", width, "dxa")


def _set_or_add_measure(parent: Any, tag: str, width: int, measure_type: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    child.set(qn("w:w"), str(width))
    child.set(qn("w:type"), measure_type)


def _shade_cell(cell: Any, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _shade_header_row(row: Any) -> None:
    for cell in row.cells:
        _shade_cell(cell, "F2F4F7")


def _set_repeat_table_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    header = properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        properties.append(header)
    header.set(qn("w:val"), "true")


def _layout_table(
    table: Table,
) -> tuple[list[list[object | TableCell | None]], list[tuple[int, int, TableCell, int, int]]]:
    rows = _safe_table_rows(table)
    row_count = len(rows)
    if not row_count:
        return [], []
    grid: list[list[object | TableCell | None]] = [[] for _ in range(row_count)]
    anchors: list[tuple[int, int, TableCell, int, int]] = []
    current_width = 0

    def ensure_width(width: int) -> None:
        nonlocal current_width
        if width > _MAX_TABLE_COLUMNS:
            raise RenderError(f"Table exceeds the safe {_MAX_TABLE_COLUMNS}-column limit")
        if width <= current_width:
            return
        for row in grid:
            row.extend([None] * (width - len(row)))
        current_width = width

    for row_index, row in enumerate(rows):
        column_index = 0
        cells = row.cells if isinstance(row.cells, list | tuple) else []
        if len(cells) > _MAX_TABLE_COLUMNS:
            raise RenderError(
                f"Table exceeds the safe {_MAX_TABLE_COLUMNS}-column limit"
            )
        for cell in cells[:_MAX_TABLE_COLUMNS]:
            if not isinstance(cell, TableCell):
                continue
            while column_index < len(grid[row_index]) and grid[row_index][column_index] is not None:
                column_index += 1
            column_span = _safe_span(cell.column_span, _MAX_TABLE_COLUMNS)
            row_span = _safe_span(cell.row_span, row_count - row_index)
            while True:
                ensure_width(column_index + column_span)
                occupied = any(
                    grid[target_row][target_column] is not None
                    for target_row in range(row_index, row_index + row_span)
                    for target_column in range(column_index, column_index + column_span)
                )
                if not occupied:
                    break
                column_index += 1
            grid[row_index][column_index] = cell
            for target_row in range(row_index, row_index + row_span):
                for target_column in range(column_index, column_index + column_span):
                    if target_row != row_index or target_column != column_index:
                        grid[target_row][target_column] = _COVERED
            anchors.append((row_index, column_index, cell, column_span, row_span))
            column_index += column_span
    max_columns = max((len(row) for row in grid), default=0)
    ensure_width(max_columns)
    return grid, anchors


def _safe_table_rows(table: Table) -> list[TableRow]:
    rows = table.rows
    if not isinstance(rows, list | tuple):
        return []
    result: list[TableRow] = []
    seen: set[int] = set()
    omitted = len(rows) > _MAX_TABLE_ROWS
    for row in rows[:_MAX_TABLE_ROWS]:
        if not isinstance(row, TableRow) or id(row) in seen:
            omitted = True
            continue
        seen.add(id(row))
        result.append(row)
    if omitted:
        result = result[: _MAX_TABLE_ROWS - 1]
        result.append(
            TableRow(
                cells=[
                    TableCell(
                        blocks=[
                            Paragraph(
                                runs=[
                                    TextRun(
                                        "[Table content omitted at safe rendering limit]"
                                    )
                                ]
                            )
                        ]
                    )
                ]
            )
        )
    return result


def _scrub_core_properties(core: Any) -> None:
    for attribute in (
        "author",
        "category",
        "comments",
        "content_status",
        "identifier",
        "keywords",
        "language",
        "last_modified_by",
        "subject",
        "title",
        "version",
    ):
        setattr(core, attribute, "")
    # python-docx requires a datetime when using these public setters.  Remove
    # the corresponding optional elements directly instead of inventing a
    # misleading date.
    for child in list(core._element):
        if child.tag.rsplit("}", 1)[-1] in {"created", "lastPrinted", "modified"}:
            core._element.remove(child)
    core.revision = 1


def _sanitize_package(payload: bytes) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(payload), "r") as source, ZipFile(output, "w") as target:
        for source_info in source.infolist():
            name = source_info.filename
            if name.startswith(("/", "../")) or "/../" in name:
                raise RenderError("DOCX generator produced an unsafe package path")
            lowered = name.casefold()
            if lowered.startswith("customxml/") or lowered.startswith(
                "docprops/thumbnail."
            ):
                continue
            if (
                lowered == "docprops/custom.xml"
                or lowered.startswith("word/embeddings/")
                or lowered.startswith("word/activex/")
                or "vbaproject" in lowered
                or "macros" in lowered
            ):
                raise RenderError(f"DOCX generator unexpectedly produced active content: {name}")
            member = source.read(source_info)
            if lowered.endswith(".rels"):
                member = _sanitize_relationships(member)
            elif lowered == "[content_types].xml":
                member = _sanitize_content_types(member)
            elif lowered == "docprops/app.xml":
                member = _sanitize_extended_properties(member)
            if lowered.startswith("word/") and lowered.endswith(".xml"):
                member = _RSID_ATTRIBUTE.sub(b"", member)
                if lowered == "word/settings.xml":
                    member = _RSID_SETTINGS.sub(b"", member)
                    member = _DOC_ID.sub(b"", member)
                    member = _SAVE_PREVIEW.sub(b"", member)
            info = ZipInfo(name, _ZIP_TIMESTAMP)
            info.compress_type = source_info.compress_type or ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = 0o600 << 16
            target.writestr(info, member)
    return output.getvalue()


def _sanitize_relationships(payload: bytes) -> bytes:
    try:
        root = ET.fromstring(payload)
    except ET.ParseError as exc:
        raise RenderError("DOCX generator produced malformed relationships XML") from exc
    namespace = root.tag.removesuffix("Relationships").strip("{}")
    if namespace:
        ET.register_namespace("", namespace)
    for relationship in list(root):
        target_mode = next(
            (
                value
                for key, value in relationship.attrib.items()
                if key.endswith("}TargetMode") or key == "TargetMode"
            ),
            "",
        )
        if target_mode.casefold() == "external":
            raise RenderError("DOCX generator unexpectedly produced an external relationship")
        relationship_type = next(
            (
                value
                for key, value in relationship.attrib.items()
                if key.endswith("}Type") or key == "Type"
            ),
            "",
        ).casefold()
        if relationship_type.endswith(("/customxml", "/metadata/thumbnail")):
            root.remove(relationship)
        elif relationship_type.endswith(
            ("/oleobject", "/package", "/attachedtemplate", "/vbaproject")
        ):
            raise RenderError("DOCX generator unexpectedly produced active content")
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _sanitize_content_types(payload: bytes) -> bytes:
    try:
        root = ET.fromstring(payload)
    except ET.ParseError as exc:
        raise RenderError("DOCX generator produced malformed content-types XML") from exc
    namespace = root.tag.removesuffix("Types").strip("{}")
    if namespace:
        ET.register_namespace("", namespace)
    for element in list(root):
        part_name = element.attrib.get("PartName", "").casefold()
        extension = element.attrib.get("Extension", "").casefold()
        if part_name.startswith(("/customxml/", "/docprops/thumbnail.")) or extension in {
            "jpeg",
            "jpg",
        }:
            root.remove(element)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _sanitize_extended_properties(payload: bytes) -> bytes:
    try:
        root = ET.fromstring(payload)
    except ET.ParseError as exc:
        raise RenderError("DOCX generator produced malformed extended-properties XML") from exc
    namespace = root.tag.removesuffix("Properties").strip("{}")
    if namespace:
        ET.register_namespace("", namespace)
    for element in root:
        local_name = element.tag.rsplit("}", 1)[-1]
        if local_name == "Application":
            element.text = "amipro-sam-toolkit"
        elif local_name in {
            "AppVersion",
            "Company",
            "HyperlinkBase",
            "Manager",
            "Template",
        }:
            element.text = ""
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _merge_character_style(base: CharacterStyle, run: CharacterStyle) -> CharacterStyle:
    return CharacterStyle(
        bold=base.bold or run.bold,
        italic=base.italic or run.italic,
        underline=base.underline or run.underline,
        strike=base.strike or run.strike,
        superscript=base.superscript or run.superscript,
        subscript=(base.subscript or run.subscript) and not (base.superscript or run.superscript),
        font_family=run.font_family or base.font_family,
        font_size_pt=run.font_size_pt or base.font_size_pt,
        color=run.color or base.color,
    )


def _resolved_style(
    document: AmiProDocument,
    name: str | None,
) -> StyleDefinition | None:
    styles = getattr(document, "styles", None)
    if not isinstance(name, str) or not name or not isinstance(styles, dict):
        return None
    if name not in styles:
        return None
    chain: list[StyleDefinition] = []
    seen: set[str] = set()
    current_name: str | None = name
    while current_name and current_name not in seen and len(chain) < 64:
        current = styles.get(current_name)
        if not isinstance(current, StyleDefinition) or not isinstance(
            current.character, CharacterStyle
        ):
            break
        chain.append(current)
        seen.add(current_name)
        current_name = current.parent if isinstance(current.parent, str) else None

    resolved = StyleDefinition(name=name)
    for item in reversed(chain):
        resolved.character = _merge_character_style(resolved.character, item.character)
        for attribute in (
            "alignment",
            "left_indent_in",
            "right_indent_in",
            "first_line_indent_in",
            "space_before_pt",
            "space_after_pt",
            "line_spacing",
        ):
            value = getattr(item, attribute)
            if value is not None:
                setattr(resolved, attribute, value)
    return resolved


def _image_placeholder(image: Image) -> str:
    alt = _safe_label_field(image.alt_text, "Embedded image", maximum=256)
    detail = f"Image: {alt}"
    if image.reference:
        reference = _safe_label_field(
            image.reference, "invalid reference omitted", maximum=256
        )
        detail += f" (source reference not opened: {reference})"
    elif image.data is not None:
        detail += " (embedded image preserved as a placeholder)"
    return f"[{detail}]"


def _sdw_placeholder(drawing: SdwDrawing) -> str:
    alt = _safe_label_field(drawing.alt_text, "Ami Draw object", maximum=256)
    status = _safe_label_field(drawing.status, "unavailable", maximum=64)
    reason = _safe_label_field(drawing.reason, "preview unavailable", maximum=256)
    details = [
        f"Ami Draw object: {alt}",
        "no valid companion preview",
        "vector payload not rendered",
        f"status={status}",
        f"reason={reason}",
    ]
    length = drawing.declared_length
    if isinstance(length, int) and not isinstance(length, bool) and 0 <= length <= 2**63 - 1:
        details.append(f"declared length={length} bytes")
    digest = drawing.source_sha256
    if isinstance(digest, str) and re.fullmatch(r"[0-9a-fA-F]{64}", digest):
        details.append(f"SHA-256={digest.lower()}")
    return "[" + "; ".join(details) + "]"


def _safe_label_field(value: object, default: str, *, maximum: int) -> str:
    if isinstance(value, str):
        result = value[: maximum * 4]
    elif isinstance(value, bytes):
        result = value[: maximum * 4].decode("utf-8", errors="replace")
    elif isinstance(value, bool | float):
        try:
            result = str(value)
        except (TypeError, ValueError, OverflowError):
            return default
    elif isinstance(value, int):
        bits = value.bit_length()
        result = (
            f"[oversized integer omitted: {bits} bits]"
            if bits > 1_024
            else str(value)
        )
    else:
        return default
    result = " ".join(result.split())[:maximum]
    return result or default


def _valid_sdw_display_size(width: object, height: object) -> bool:
    return (
        isinstance(width, int | float)
        and not isinstance(width, bool)
        and isinstance(height, int | float)
        and not isinstance(height, bool)
        and math.isfinite(width)
        and math.isfinite(height)
        and 0.0 < width <= 6.25
        and 0.0 < height <= 7.5
    )


def _clean_xml_text(text: str) -> str:
    return "".join(
        character if _is_xml_character(ord(character)) else "\ufffd"
        for character in text
    )


def _is_xml_character(codepoint: int) -> bool:
    return (
        codepoint in {0x9, 0xA, 0xD}
        or 0x20 <= codepoint <= 0xD7FF
        or 0xE000 <= codepoint <= 0xFFFD
        or 0x10000 <= codepoint <= 0x10FFFF
    )


def _color(value: str | None) -> str | None:
    if not isinstance(value, str) or not value:
        return None
    match = _HEX_COLOR.fullmatch(value.strip())
    return match.group(1).upper() if match else None


def _number(value: float | None, default: float, minimum: float, maximum: float) -> float:
    try:
        number = float(value) if value is not None else default
    except (TypeError, ValueError, OverflowError):
        return default
    if not math.isfinite(number):
        return default
    return min(maximum, max(minimum, number))


def _first_not_none(first: float | None, second: float | None) -> float | None:
    return first if first is not None else second


def _safe_span(value: object, maximum: int) -> int:
    try:
        return max(1, min(int(value), maximum))
    except (TypeError, ValueError, OverflowError):
        return 1


def _safe_level(value: object) -> int:
    try:
        return max(0, min(int(value), 15))
    except (TypeError, ValueError, OverflowError):
        return 0
